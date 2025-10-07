#!/usr/bin/env python3
"""
AI-POWERED DOCUMENT REPORT GENERATOR WITH ADVANCED OCR
=======================================================
Automatically analyzes ANY document including scanned images with OCR
and generates intelligent insights using LLM.

Features:
- Advanced OCR with image preprocessing
- Multi-format document parsing (PDF, Excel, CSV, TXT, DOCX, Images)
- AI-powered analysis and insights
- Automated chart generation
- Professional PDF reports

Requirements:
pip install pandas matplotlib reportlab openpyxl PyPDF2 python-docx anthropic pytesseract pillow pdf2image opencv-python
"""

import os
import sys
import argparse
from datetime import datetime
from pathlib import Path
import json
import re

# Check required libraries
try:
    import pandas as pd
    import matplotlib.pyplot as plt
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
    print("✅ Core libraries loaded")
except ImportError as e:
    print(f"❌ Missing core library: {e}")
    print("Install: pip install pandas matplotlib reportlab")
    sys.exit(1)

# Optional libraries for document parsing
OPTIONAL_IMPORTS = {}

try:
    import PyPDF2
    OPTIONAL_IMPORTS['pdf'] = True
except:
    OPTIONAL_IMPORTS['pdf'] = False
    print("⚠️  PyPDF2 not available (pip install PyPDF2)")

try:
    import docx
    OPTIONAL_IMPORTS['docx'] = True
except:
    OPTIONAL_IMPORTS['docx'] = False
    print("⚠️  python-docx not available (pip install python-docx)")

try:
    import anthropic
    OPTIONAL_IMPORTS['anthropic'] = True
except:
    OPTIONAL_IMPORTS['anthropic'] = False
    print("⚠️  anthropic not available (pip install anthropic)")

# OCR libraries
try:
    import pytesseract
    from PIL import Image as PILImage
    import cv2
    import numpy as np
    OPTIONAL_IMPORTS['ocr'] = True
    print("✅ OCR libraries loaded (Tesseract + OpenCV)")
except ImportError as e:
    OPTIONAL_IMPORTS['ocr'] = False
    print(f"⚠️  OCR not available: {e}")
    print("   Install: pip install pytesseract pillow opencv-python")
    print("   Also install Tesseract-OCR:")
    print("   - Windows: https://github.com/UB-Mannheim/tesseract/wiki")
    print("   - Mac: brew install tesseract")
    print("   - Linux: sudo apt-get install tesseract-ocr")

try:
    from pdf2image import convert_from_path
    OPTIONAL_IMPORTS['pdf2image'] = True
    print("✅ PDF to Image conversion available")
except:
    OPTIONAL_IMPORTS['pdf2image'] = False
    print("⚠️  pdf2image not available (pip install pdf2image)")


class OCRProcessor:
    """Advanced OCR processor with image preprocessing"""
    
    def __init__(self):
        if not OPTIONAL_IMPORTS.get('ocr'):
            raise ImportError("OCR libraries not available. Install: pip install pytesseract pillow opencv-python")
        
        # Try to find Tesseract executable
        self._setup_tesseract()
    
    def _setup_tesseract(self):
        """Setup Tesseract path for different OS"""
        # Common Tesseract paths
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            '/usr/local/bin/tesseract',
            '/usr/bin/tesseract',
            '/opt/homebrew/bin/tesseract'
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"✅ Tesseract found at: {path}")
                return
        
        # Try to use system PATH
        try:
            pytesseract.get_tesseract_version()
            print("✅ Tesseract found in system PATH")
        except:
            print("⚠️  Tesseract not found. Please install Tesseract-OCR")
            print("   Download from: https://github.com/tesseract-ocr/tesseract")
    
    def preprocess_image(self, image_path, output_path=None):
        """
        Advanced image preprocessing for better OCR accuracy
        
        Techniques:
        1. Grayscale conversion
        2. Noise removal
        3. Thresholding
        4. Deskewing
        5. Border removal
        """
        print(f"  🔧 Preprocessing image: {Path(image_path).name}")
        
        # Read image
        img = cv2.imread(str(image_path))
        if img is None:
            raise ValueError(f"Could not read image: {image_path}")
        
        # 1. Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # 2. Noise removal (denoising)
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        
        # 3. Thresholding - multiple methods, choose best
        # Method A: Otsu's thresholding
        _, thresh1 = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Method B: Adaptive thresholding
        thresh2 = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY, 11, 2
        )
        
        # Method C: Simple binary threshold
        _, thresh3 = cv2.threshold(denoised, 127, 255, cv2.THRESH_BINARY)
        
        # Use Otsu's method as default (usually best)
        thresh = thresh1
        
        # 4. Deskew (straighten tilted images)
        coords = np.column_stack(np.where(thresh > 0))
        if len(coords) > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            
            if abs(angle) > 0.5:  # Only deskew if significant tilt
                (h, w) = thresh.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                thresh = cv2.warpAffine(
                    thresh, M, (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE
                )
                print(f"    ↻ Deskewed by {angle:.2f}°")
        
        # 5. Remove borders
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        if contours:
            # Find largest contour (usually the document)
            largest_contour = max(contours, key=cv2.contourArea)
            x, y, w, h = cv2.boundingRect(largest_contour)
            
            # Add small margin
            margin = 10
            x = max(0, x - margin)
            y = max(0, y - margin)
            w = min(thresh.shape[1] - x, w + 2 * margin)
            h = min(thresh.shape[0] - y, h + 2 * margin)
            
            thresh = thresh[y:y+h, x:x+w]
        
        # 6. Enhance contrast
        thresh = cv2.equalizeHist(thresh)
        
        # 7. Morphological operations to clean up
        kernel = np.ones((1, 1), np.uint8)
        thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        thresh = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
        
        # Save preprocessed image if path provided
        if output_path:
            cv2.imwrite(str(output_path), thresh)
            print(f"    💾 Saved preprocessed: {Path(output_path).name}")
        
        return thresh
    
    def extract_text_from_image(self, image_path, preprocess=True, lang='eng'):
        """
        Extract text from image using OCR
        
        Args:
            image_path: Path to image file
            preprocess: Whether to preprocess image (recommended)
            lang: Tesseract language (eng, fra, deu, spa, etc.)
        """
        try:
            if preprocess:
                # Preprocess image
                processed_img = self.preprocess_image(image_path)
                
                # Convert numpy array to PIL Image
                pil_img = PILImage.fromarray(processed_img)
            else:
                # Use original image
                pil_img = PILImage.open(image_path)
            
            # Configure Tesseract for better accuracy
            custom_config = r'--oem 3 --psm 6'
            # OEM 3 = LSTM neural network mode
            # PSM 6 = Assume uniform block of text
            
            print(f"  🔍 Running OCR on: {Path(image_path).name}")
            
            # Extract text
            text = pytesseract.image_to_string(pil_img, lang=lang, config=custom_config)
            
            # Also get detailed data for confidence scores
            data = pytesseract.image_to_data(pil_img, lang=lang, config=custom_config, output_type=pytesseract.Output.DICT)
            
            # Calculate average confidence
            confidences = [int(conf) for conf in data['conf'] if conf != '-1']
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            print(f"    ✅ OCR complete - Confidence: {avg_confidence:.1f}%")
            
            # Clean up text
            text = text.strip()
            
            return {
                'text': text,
                'confidence': avg_confidence,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
            
        except Exception as e:
            print(f"    ❌ OCR failed: {e}")
            return {
                'text': '',
                'confidence': 0,
                'error': str(e)
            }
    
    def extract_text_from_pdf_with_ocr(self, pdf_path):
        """Extract text from PDF using OCR (for scanned PDFs)"""
        if not OPTIONAL_IMPORTS.get('pdf2image'):
            return {'error': 'pdf2image not installed. Run: pip install pdf2image'}
        
        try:
            print(f"  📄 Converting PDF pages to images...")
            
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)
            
            all_text = []
            total_confidence = 0
            
            for i, img in enumerate(images, 1):
                print(f"  📄 Processing page {i}/{len(images)}")
                
                # Save temporary image
                temp_img_path = f"temp_page_{i}.png"
                img.save(temp_img_path, 'PNG')
                
                # OCR the page
                result = self.extract_text_from_image(temp_img_path)
                all_text.append(f"\n--- Page {i} ---\n{result['text']}")
                total_confidence += result.get('confidence', 0)
                
                # Clean up temp file
                os.remove(temp_img_path)
            
            combined_text = '\n'.join(all_text)
            avg_confidence = total_confidence / len(images) if images else 0
            
            return {
                'text': combined_text,
                'pages': len(images),
                'confidence': avg_confidence,
                'word_count': len(combined_text.split()),
                'char_count': len(combined_text)
            }
            
        except Exception as e:
            print(f"    ❌ PDF OCR failed: {e}")
            return {'error': str(e)}


class DocumentParser:
    """Parse various document formats with OCR support"""
    
    def __init__(self):
        self.ocr_processor = None
        if OPTIONAL_IMPORTS.get('ocr'):
            try:
                self.ocr_processor = OCRProcessor()
            except Exception as e:
                print(f"⚠️  OCR initialization failed: {e}")
    
    def parse_csv(self, file_path):
        """Parse CSV file"""
        try:
            df = pd.read_csv(file_path)
            return {
                'type': 'csv',
                'dataframe': df,
                'text': df.to_string(),
                'summary': f"CSV with {len(df)} rows and {len(df.columns)} columns",
                'columns': list(df.columns)
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_excel(self, file_path):
        """Parse Excel file"""
        try:
            df = pd.read_excel(file_path)
            return {
                'type': 'excel',
                'dataframe': df,
                'text': df.to_string(),
                'summary': f"Excel with {len(df)} rows and {len(df.columns)} columns",
                'columns': list(df.columns)
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_pdf(self, file_path):
        """Parse PDF file with OCR fallback"""
        if not OPTIONAL_IMPORTS.get('pdf'):
            return {'error': 'PyPDF2 not installed. Run: pip install PyPDF2'}
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Try text extraction first
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            # If very little text extracted, try OCR
            if len(text.strip()) < 100 and self.ocr_processor:
                print("  ℹ️  Limited text found, attempting OCR...")
                ocr_result = self.ocr_processor.extract_text_from_pdf_with_ocr(file_path)
                
                if 'error' not in ocr_result:
                    text = ocr_result['text']
                    return {
                        'type': 'pdf_ocr',
                        'text': text,
                        'summary': f"PDF with {ocr_result.get('pages', 0)} pages (OCR processed)",
                        'confidence': ocr_result.get('confidence', 0),
                        'method': 'OCR'
                    }
            
            return {
                'type': 'pdf',
                'text': text,
                'summary': f"PDF with {len(reader.pages)} pages",
                'method': 'Text Extraction'
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_image(self, file_path):
        """Parse image file using OCR"""
        if not self.ocr_processor:
            return {'error': 'OCR not available. Install: pip install pytesseract pillow opencv-python'}
        
        try:
            result = self.ocr_processor.extract_text_from_image(file_path)
            
            if 'error' in result:
                return result
            
            return {
                'type': 'image_ocr',
                'text': result['text'],
                'summary': f"Image with {result['word_count']} words extracted (confidence: {result['confidence']:.1f}%)",
                'confidence': result['confidence'],
                'method': 'OCR'
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_docx(self, file_path):
        """Parse DOCX file"""
        if not OPTIONAL_IMPORTS.get('docx'):
            return {'error': 'python-docx not installed. Run: pip install python-docx'}
        
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                'type': 'docx',
                'text': text,
                'summary': f"Word document with {len(doc.paragraphs)} paragraphs"
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_txt(self, file_path):
        """Parse text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'type': 'txt',
                'text': text,
                'summary': f"Text file with {len(text)} characters"
            }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_document(self, file_path):
        """Auto-detect and parse document"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {'error': f'File not found: {file_path}'}
        
        ext = file_path.suffix.lower()
        
        # Image formats (OCR)
        image_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}
        
        if ext in image_formats:
            return self.parse_image(file_path)
        
        # Other formats
        parsers = {
            '.csv': self.parse_csv,
            '.xlsx': self.parse_excel,
            '.xls': self.parse_excel,
            '.pdf': self.parse_pdf,
            '.docx': self.parse_docx,
            '.doc': self.parse_docx,
            '.txt': self.parse_txt,
        }
        
        parser = parsers.get(ext)
        if not parser:
            return {'error': f'Unsupported file type: {ext}'}
        
        return parser(file_path)


class AIAnalyzer:
    """AI-powered document analyzer using Claude"""
    
    def __init__(self, api_key=None):
        self.api_key = api_key or os.environ.get('ANTHROPIC_API_KEY')
        
        if not self.api_key:
            print("⚠️  No API key found. Set ANTHROPIC_API_KEY environment variable.")
            print("   Get your key from: https://console.anthropic.com/")
            self.client = None
        else:
            if not OPTIONAL_IMPORTS.get('anthropic'):
                print("❌ Anthropic library not installed. Run: pip install anthropic")
                self.client = None
            else:
                self.client = anthropic.Anthropic(api_key=self.api_key)
                print("✅ AI analyzer ready")
    
    def analyze_document(self, parsed_doc, custom_prompt=None):
        """Analyze document using Claude"""
        
        if not self.client:
            return self._fallback_analysis(parsed_doc)
        
        # Prepare document content
        doc_text = parsed_doc.get('text', '')[:50000]  # Limit to ~50k chars
        doc_type = parsed_doc.get('type', 'unknown')
        doc_summary = parsed_doc.get('summary', '')
        
        # Add OCR confidence info if available
        confidence_info = ""
        if 'confidence' in parsed_doc:
            confidence_info = f"\nOCR Confidence: {parsed_doc['confidence']:.1f}%"
        
        # Build analysis prompt
        base_prompt = f"""You are an expert data analyst. Analyze this {doc_type} document and provide comprehensive insights.

Document Info: {doc_summary}{confidence_info}

Document Content:
{doc_text}

Please provide a comprehensive analysis in JSON format with these sections:

1. executive_summary: A concise 2-3 sentence overview of the document
2. key_findings: List of 5-7 most important insights (as array of strings)
3. data_insights: Detailed analysis of patterns, trends, and anomalies
4. recommendations: 5-7 actionable recommendations (as array of strings)
5. metrics: Key numbers and statistics found (as object with metric names and values)
6. risks_concerns: Potential issues or risks identified (as array of strings)
7. opportunities: Growth or improvement opportunities (as array of strings)

{custom_prompt or ''}

Return ONLY valid JSON, no other text."""

        try:
            print("\n🤖 Analyzing with Claude AI...")
            
            message = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                temperature=0.7,
                messages=[
                    {"role": "user", "content": base_prompt}
                ]
            )
            
            response_text = message.content[0].text
            
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
                print("✅ AI analysis complete")
                return analysis
            else:
                return json.loads(response_text)
                
        except Exception as e:
            print(f"⚠️  AI analysis failed: {e}")
            return self._fallback_analysis(parsed_doc)
    
    def _fallback_analysis(self, parsed_doc):
        """Fallback analysis without AI"""
        return {
            'executive_summary': f"This document contains {parsed_doc.get('summary', 'data')}.",
            'key_findings': [
                'Document successfully parsed',
                'Manual review recommended for detailed insights',
                'Consider adding AI analysis for deeper insights'
            ],
            'data_insights': 'Basic parsing completed. Enable AI analysis for detailed insights.',
            'recommendations': [
                'Review the document content carefully',
                'Set up ANTHROPIC_API_KEY for AI-powered analysis',
                'Consider exporting data for further analysis'
            ],
            'metrics': {},
            'risks_concerns': ['Limited analysis without AI'],
            'opportunities': ['Enable AI for comprehensive insights']
        }


class ReportGenerator:
    """Generate comprehensive reports with charts and insights"""
    
    def __init__(self, output_dir="ai_reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        print(f"📁 Output directory: {self.output_dir.absolute()}")
    
    def create_visualizations(self, parsed_doc):
        """Create charts from data"""
        chart_paths = []
        
        if 'dataframe' not in parsed_doc:
            print("ℹ️  No tabular data for visualizations")
            return chart_paths
        
        df = parsed_doc['dataframe']
        print(f"\n📊 Generating visualizations from {len(df)} rows...")
        
        # Try to create meaningful charts
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        
        if len(numeric_cols) == 0:
            print("  ⚠️  No numeric columns found")
            return chart_paths
        
        # Chart 1: Overview of numeric columns
        if len(numeric_cols) > 0:
            plt.figure(figsize=(12, 6))
            
            if len(df) <= 20:  # Bar chart for small datasets
                df[numeric_cols[:5]].plot(kind='bar', figsize=(12, 6))
                plt.title('Data Overview (Bar Chart)', fontsize=14, fontweight='bold')
            else:  # Line chart for larger datasets
                df[numeric_cols[:5]].plot(figsize=(12, 6))
                plt.title('Data Overview (Trends)', fontsize=14, fontweight='bold')
            
            plt.xlabel('Index/Category')
            plt.ylabel('Values')
            plt.legend(loc='best')
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
            
            chart_path = self.output_dir / "chart_overview.png"
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_path))
            print(f"  ✅ Created overview chart")
        
        # Chart 2: Distribution/Statistics
        if len(numeric_cols) >= 2:
            fig, axes = plt.subplots(1, min(3, len(numeric_cols)), figsize=(15, 5))
            if len(numeric_cols) == 2:
                axes = [axes]
            
            for idx, col in enumerate(numeric_cols[:3]):
                ax = axes[idx] if len(numeric_cols) > 2 else axes[0]
                df[col].hist(bins=20, ax=ax, color='steelblue', alpha=0.7)
                ax.set_title(f'{col} Distribution')
                ax.set_xlabel('Value')
                ax.set_ylabel('Frequency')
                ax.grid(True, alpha=0.3)
            
            plt.tight_layout()
            chart_path = self.output_dir / "chart_distribution.png"
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_path))
            print(f"  ✅ Created distribution chart")
        
        return chart_paths
    
    def generate_pdf_report(self, parsed_doc, analysis, chart_paths, document_name):
        """Create professional PDF report"""
        print("\n📄 Generating PDF report...")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.output_dir / f"report_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=36
        )
        
        content = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY
        ))
        
        # Title
        content.append(Paragraph("AI-Powered Document Analysis Report", styles['Title']))
        content.append(Spacer(1, 0.3*inch))
        
        # Metadata
        method_info = parsed_doc.get('method', 'Unknown')
        confidence_info = ""
        if 'confidence' in parsed_doc:
            confidence_info = f"<br/><b>OCR Confidence:</b> {parsed_doc['confidence']:.1f}%"
        
        meta_text = f"""
        <b>Document:</b> {document_name}<br/>
        <b>Type:</b> {parsed_doc.get('type', 'unknown').upper()}<br/>
        <b>Processing Method:</b> {method_info}{confidence_info}<br/>
        <b>Analysis Date:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}<br/>
        <b>Summary:</b> {parsed_doc.get('summary', 'N/A')}
        """
        content.append(Paragraph(meta_text, styles['Normal']))
        content.append(Spacer(1, 0.4*inch))
        
        # Executive Summary
        content.append(Paragraph("Executive Summary", styles['Heading1']))
        content.append(Spacer(1, 12))
        content.append(Paragraph(analysis.get('executive_summary', 'N/A'), styles['CustomBody']))
        content.append(Spacer(1, 0.3*inch))
        
        # Key Findings
        content.append(Paragraph("Key Findings", styles['Heading1']))
        content.append(Spacer(1, 12))
        findings = analysis.get('key_findings', [])
        for finding in findings:
            content.append(Paragraph(f"• {finding}", styles['CustomBody']))
            content.append(Spacer(1, 6))
        content.append(Spacer(1, 0.3*inch))
        
        # Data Insights
        if analysis.get('data_insights'):
            content.append(Paragraph("Detailed Insights", styles['Heading1']))
            content.append(Spacer(1, 12))
            content.append(Paragraph(analysis['data_insights'], styles['CustomBody']))
            content.append(Spacer(1, 0.3*inch))
        
        # Metrics
        if analysis.get('metrics'):
            content.append(Paragraph("Key Metrics", styles['Heading1']))
            content.append(Spacer(1, 12))
            for metric, value in analysis['metrics'].items():
                content.append(Paragraph(f"<b>{metric}:</b> {value}", styles['Normal']))
                content.append(Spacer(1, 6))
            content.append(Spacer(1, 0.3*inch))
        
        # Visualizations
        if chart_paths:
            content.append(PageBreak())
            content.append(Paragraph("Data Visualizations", styles['Heading1']))
            content.append(Spacer(1, 12))
            
            for chart_path in chart_paths:
                try:
                    img = Image(chart_path, width=6*inch, height=3*inch)
                    content.append(img)
                    content.append(Spacer(1, 0.2*inch))
                except Exception as e:
                    print(f"  ⚠️  Could not add chart: {e}")
        
        # Recommendations
        content.append(PageBreak())
        content.append(Paragraph("Recommendations", styles['Heading1']))
        content.append(Spacer(1, 12))
        recommendations = analysis.get('recommendations', [])
        for i, rec in enumerate(recommendations, 1):
            content.append(Paragraph(f"<b>{i}.</b> {rec}", styles['CustomBody']))
            content.append(Spacer(1, 10))
        
        # Risks & Opportunities
        if analysis.get('risks_concerns'):
            content.append(Spacer(1, 0.2*inch))
            content.append(Paragraph("Risks & Concerns", styles['Heading2']))
            content.append(Spacer(1, 8))
            for risk in analysis['risks_concerns']:
                content.append(Paragraph(f"⚠️ {risk}", styles['Normal']))
                content.append(Spacer(1, 6))

        # References
        if analysis.get('references'):
            content.append(PageBreak())
            content.append(Paragraph("References", styles['Heading1']))
            content.append(Spacer(1, 12))
            for i, ref in enumerate(analysis['references'], 1):
                content.append(Paragraph(f"[{i}] {ref}", styles['Normal']))
                content.append(Spacer(1, 6))

        # Build PDF
        try:
            doc.build(content)
            print(f"✅ PDF saved: {output_path}")
            return str(output_path)
        except Exception as e:
            print(f"❌ PDF build error: {e}")
            return ""


def main():
    """CLI entrypoint: parse documents, analyze with AI, generate PDF report(s)."""
    parser = argparse.ArgumentParser(description="AI Document Report Generator (OCR + AI + Charts)")
    parser.add_argument("--input", "-i", required=True, help="Path to input file or directory")
    parser.add_argument("--output", "-o", default="ai_reports", help="Output directory for reports and charts")
    parser.add_argument("--prompt", "-p", default=None, help="Additional instructions for the AI analyzer")
    parser.add_argument("--no-charts", action="store_true", help="Skip chart generation even if dataframes exist")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Input path not found: {input_path}")
        sys.exit(1)

    doc_parser = DocumentParser()
    analyzer = AIAnalyzer()
    reporter = ReportGenerator(output_dir=args.output)

    def process_file(file_path: Path):
        print(f"\n===== Processing: {file_path.name} =====")
        parsed = doc_parser.parse_document(file_path)
        if 'error' in parsed:
            print(f"❌ Parse error: {parsed['error']}")
            return
        analysis = analyzer.analyze_document(parsed, custom_prompt=args.prompt)
        charts = [] if args.no_charts else reporter.create_visualizations(parsed)
        report_path = reporter.generate_pdf_report(parsed, analysis, charts, document_name=file_path.name)
        if report_path:
            print(f"✅ Report generated: {report_path}")

    if input_path.is_dir():
        supported_exts = {".pdf", ".docx", ".doc", ".csv", ".xlsx", ".xls", ".txt", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif"}
        files = [p for p in input_path.iterdir() if p.is_file() and p.suffix.lower() in supported_exts]
        if not files:
            print("ℹ️  No supported files found in directory.")
            sys.exit(0)
        for f in files:
            process_file(f)
    else:
        process_file(input_path)


if __name__ == "__main__":
    main()